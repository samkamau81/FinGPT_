import streamlit as st
from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field, validator, ValidationError
from typing import List, Optional
import asyncio
from langchain.llms import OpenAI
import docx
import datetime
import time
import base64
import os
import uvicorn
import threading
import requests

class StartupDetails(BaseModel):
    """Validated model for startup information"""
    name: str = Field(..., min_length=1, max_length=200)
    description: str = Field(..., min_length=10, max_length=2000)
    country: str
    sectors: List[str] = Field(..., min_items=1, max_items=10)
    funding_types: List[str] = Field(..., min_items=1, max_items=10)
    
    @validator('name', 'description')
    def strip_whitespace(cls, v):
        return v.strip()
    
    @validator('sectors', 'funding_types')
    def validate_not_empty(cls, v):
        if not v:
            raise ValueError('Must select at least one option')
        return v

class ReportRequest(BaseModel):
    """Request model for report generation"""
    startup_details: StartupDetails
    openai_api_key: str = Field(..., regex=r'^sk-.*')

class ChatRequest(BaseModel):
    """Request model for chat interactions"""
    message: str = Field(..., min_length=1, max_length=5000)
    openai_api_key: str = Field(..., regex=r'^sk-.*')

class ChatResponse(BaseModel):
    """Response model for chat"""
    response: str
    timestamp: datetime.datetime

class ReportResponse(BaseModel):
    """Response model for report generation"""
    success: bool
    message: str
    report_path: Optional[str] = None

#FASTAPI BACKEND
app = FastAPI(title="MoneyMentor API", version="1.0.0")

async def generate_llm_response_async(prompt: str, api_key: str) -> str:
    """Generate LLM response asynchronously"""
    try:
        llm = OpenAI(temperature=0.3, openai_api_key=api_key)
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(None, llm, prompt)
        return response
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM error: {str(e)}")

def generate_report_doc(company_name: str, country: str, funding_summary: str, 
                        legal_summary: str) -> str:
    """Generate Word document report"""
    doc = docx.Document()
    report_date = datetime.date.today()
    
    doc.add_heading("Financial Report", 0)
    doc.add_paragraph('Authored By: MoneyMentor FinGPT LLM')
    doc.add_paragraph(f'Created On: {str(report_date)}')
    doc.add_paragraph(f'Created For: {company_name}')
    doc.add_paragraph(f'Country based: {country}')
    doc.add_heading(f'Navigating the Intersection: A Comprehensive Guide to {company_name} Financial and Legal Strategies')
    
    doc.add_heading('Funding Strategies')
    doc.add_paragraph(funding_summary)
    
    doc.add_heading('Legal Strategies')
    doc.add_paragraph(legal_summary)
    
    os.makedirs("reports", exist_ok=True)
    report_path = f"reports/Financial_Report_{company_name.replace(' ', '_')}_{report_date}.docx"
    doc.save(report_path)
    
    return report_path

@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(request: ChatRequest):
    """Handle chat requests with validation"""
    try:
        response = await generate_llm_response_async(request.message, request.openai_api_key)
        return ChatResponse(
            response=response,
            timestamp=datetime.datetime.now()
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate-report", response_model=ReportResponse)
async def generate_report_endpoint(request: ReportRequest):
    """Generate comprehensive financial report"""
    startup = request.startup_details
    
    try:
        # Funding prompt
        funding_prompt = f"""I'm exploring funding options for my startup named {startup.name}.
Description: {startup.description}
I'm interested in {', '.join(startup.funding_types)} funding types.
My startup is in {startup.country} and operates in {', '.join(startup.sectors)}.

Provide comprehensive overview of:
1. Available funding sources for early-stage startups in {startup.country}
2. Eligibility requirements for {', '.join(startup.funding_types)}
3. Application/pitch preparation tips
4. Networking strategies with investors
5. Common challenges and how to avoid them
6. Timing and planning considerations

Format in detailed bullet points with up-to-date information."""

        funding_summary = await generate_llm_response_async(funding_prompt, request.openai_api_key)
        
        # Legal prompt
        legal_prompt = f"""Provide comprehensive legal guidance for launching {startup.name} in {startup.country}.
Sectors: {', '.join(startup.sectors)}

Cover all relevant legal requirements including:
1. Business registration requirements
2. Regulatory compliance for {', '.join(startup.sectors)}
3. Intellectual property protection
4. Employment law considerations
5. Tax obligations and incentives
6. Industry-specific regulations
7. Data protection and privacy laws

Format in detailed bullet points with current, country-specific information."""

        legal_summary = await generate_llm_response_async(legal_prompt, request.openai_api_key)
        
        # Generate document
        report_path = generate_report_doc(
            startup.name,
            startup.country,
            funding_summary,
            legal_summary
        )
        
        return ReportResponse(
            success=True,
            message="Report generated successfully",
            report_path=report_path
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")

@app.get("/api/health")
async def health_check():
    """API health check"""
    return {"status": "healthy", "timestamp": datetime.datetime.now().isoformat()}

# STREAMLIT FRONTEND

def run_fastapi():
    """Run FastAPI server in background thread"""
    uvicorn.run(app, host="127.0.0.1", port=8000, log_level="error")

def start_backend():
    """Start FastAPI backend if not already running"""
    try:
        response = requests.get("http://127.0.0.1:8000/api/health", timeout=1)
        if response.status_code == 200:
            return True
    except:
        thread = threading.Thread(target=run_fastapi, daemon=True)
        thread.start()
        time.sleep(2)  # Give server time to start
        return True

def call_chat_api(message: str, api_key: str) -> str:
    """Call FastAPI chat endpoint"""
    try:
        response = requests.post(
            "http://127.0.0.1:8000/api/chat",
            json={"message": message, "openai_api_key": api_key},
            timeout=120
        )
        if response.status_code == 200:
            return response.json()["response"]
        else:
            return f"Error: {response.json().get('detail', 'Unknown error')}"
    except requests.exceptions.RequestException as e:
        return f"Connection error: {str(e)}"

def call_report_api(startup_data: dict, api_key: str) -> dict:
    """Call FastAPI report generation endpoint"""
    try:
        response = requests.post(
            "http://127.0.0.1:8000/api/generate-report",
            json={
                "startup_details": startup_data,
                "openai_api_key": api_key
            },
            timeout=300
        )
        if response.status_code == 200:
            return response.json()
        else:
            return {"success": False, "message": response.json().get('detail', 'Unknown error')}
    except requests.exceptions.RequestException as e:
        return {"success": False, "message": f"Connection error: {str(e)}"}

# Initialize backend
start_backend()

# Streamlit UI
st.set_page_config(page_title="MoneyMentor", page_icon="💸", layout="wide")

# Sidebar
with st.sidebar:
    st.markdown("<h1 style='text-align:center;font-family:Georgia'>💸 MoneyMentor </h1>", unsafe_allow_html=True)
    st.markdown("""This app is designed to assist Entrepreneurs and Startup founders in navigating the financial and legal aspects of 
                developing their businesses. It provides valuable information and tools related to startup financing, addressing 
                various Legal considerations such as Locale and Type of Funding Sought using the power of Large Language Models by OpenAI and LangChain.""")
    
    st.markdown("<h2 style='text-align:center;font-family:Georgia'>Features</h2>", unsafe_allow_html=True)
    st.markdown("- 🤑 MoneyMentor FinanceGPT - AI-powered business advisor")
    st.markdown("- 🧾 Financial Report Generator - Generate comprehensive reports")
    st.markdown("- ✅ Input Validation - Ensures data quality with Pydantic")
    st.markdown("- ⚡ FastAPI Backend - Scalable API architecture")
    st.markdown("-------")
    
    openai_api_key = st.text_input('Enter OpenAI API Key', type='password')
    
    st.markdown("-------")
    st.markdown("<h1 style='text-align:center;font-family:Georgia'>🧾 Financial Report Generator</h1>", unsafe_allow_html=True)
    
    start_up_name = st.text_input("What is the name of your Start Up?")
    start_up_description = st.text_area("Describe your startup and revenue model", height=100)
    
    country = st.selectbox(
        'Where is your Start Up based?',
        ("United States of America", "United Kingdom", "Canada", "Germany", "France", "India", "China", 
         "Australia", "Singapore", "Kenya", "Nigeria", "South Africa", "Brazil", "Mexico", "Japan")
    )
    
    sector = st.multiselect(
        'What is your Start Up about?',
        ["Technology and Software", "Healthcare and Biotechnology", "E-commerce and Retail", 
         "Fintech (Financial Technology)", "Agriculture and AgriTech", "Clean Energy and Sustainability",
         "Education and EdTech", "AI and Machine Learning", "Cybersecurity", "Blockchain and Cryptocurrency"]
    )
    
    funding = st.multiselect(
        'What sort of Funding are you looking for?',
        ["Angel Investment", "Venture Capital", "Seed Funding", "Series A Funding", 
         "Crowdfunding", "Government Grants", "Accelerator Programs", "Corporate Investment"]
    )
    
    st.markdown("-------")
    generate_button = st.button("Generate Financial Report", type="primary")

# Main Chat Interface
st.markdown("<h1 style='text-align:center;font-family:Georgia'>🤑 MoneyMentor Chatbot</h1>", unsafe_allow_html=True)

# Initialize chat history
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({
        "role": "assistant",
        "content": "Hey there! I'm your MoneyMentor, here to advise you on how to kickstart your Business. Ask me anything, but first, enter your API Key in the sidebar."
    })

# Display chat history
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Handle report generation
if generate_button:
    if not openai_api_key.startswith('sk-'):
        st.warning('Please enter your OpenAI API key!', icon='⚠')
    else:
        # Validate inputs
        validation_errors = []
        if not start_up_name or len(start_up_name.strip()) < 1:
            validation_errors.append("Startup name is required")
        if not start_up_description or len(start_up_description.strip()) < 10:
            validation_errors.append("Description must be at least 10 characters")
        if not sector:
            validation_errors.append("Please select at least one sector")
        if not funding:
            validation_errors.append("Please select at least one funding type")
        
        if validation_errors:
            st.error("Validation Errors:\n" + "\n".join(f"- {err}" for err in validation_errors))
        else:
            with st.spinner('Generating comprehensive financial report...'):
                startup_data = {
                    "name": start_up_name.strip(),
                    "description": start_up_description.strip(),
                    "country": country,
                    "sectors": sector,
                    "funding_types": funding
                }
                
                result = call_report_api(startup_data, openai_api_key)
                
                if result.get("success"):
                    st.success(result["message"])
                    report_path = result.get("report_path")
                    
                    if report_path and os.path.exists(report_path):
                        with open(report_path, "rb") as file:
                            st.download_button(
                                label="📥 Download Financial Report",
                                data=file,
                                file_name=os.path.basename(report_path),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                else:
                    st.error(f"Report generation failed: {result.get('message')}")

# Handle chat input
if prompt := st.chat_input("Ask me anything about your startup..."):
    if not openai_api_key.startswith('sk-'):
        st.warning('Please enter your OpenAI API key!', icon='⚠')
    else:
        # Display user message
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        # Generate and display assistant response
        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            full_response = ""
            
            with st.spinner('Thinking...'):
                assistant_response = call_chat_api(prompt, openai_api_key)
            
            # Simulate typing effect
            for chunk in assistant_response.split():
                full_response += chunk + " "
                time.sleep(0.05)
                message_placeholder.markdown(full_response + "▌")
            
            message_placeholder.markdown(full_response)
        
        st.session_state.messages.append({"role": "assistant", "content": full_response})

st.markdown("---")
st.markdown("**💡 Tip:** Use specific questions for better responses. Example: 'What are the key legal requirements for a fintech startup in Kenya?'")
